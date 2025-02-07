import azure.functions as func
import logging
import json
import re
import os,openai
from azure.storage.blob import BlobServiceClient
from PyPDF2 import PdfReader
from docx import Document
from langchain_openai import AzureOpenAIEmbeddings  #AZure Search 
from langchain_community.vectorstores import AzureSearch # Vectors
from dotenv import load_dotenv
from nltk.tokenize import sent_tokenize #Natural Language Toolkit--->spliting purpose 
import nltk  #analyses the human data
from openai import AzureOpenAI
import io #bytes

# Ensure NLTK punkt data is downloaded
nltk.download('punkt') #split txt into smallr chunkes 

# Load environment variables from the .env file
load_dotenv()

# Log loaded environment variables
logging.info(f"Loaded Environment Variables: {os.environ.get('BLOB_CONNECTION_STRING')}, {os.environ.get('AZURE_OPENAI_API_KEY')}, {os.environ.get('AZURE_SEARCH_ENDPOINT')}")

# Log the OpenAI API Key (for debugging purposes)
logging.info("Using Azure OpenAI API Key.")

# Load environment variables for Azure credentials
BLOB_CONNECTION_STRING = os.getenv("BLOB_CONNECTION_STRING")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_API_KEY = os.getenv("AZURE_SEARCH_API_KEY")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_INDEX_NAME")
SYSTEM_MESSAGE = os.getenv("SYSTEM_MESSAGE")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")

# Specify the API version you're using
openai.api_version = "2024-10-21"  # Ensure this is the correct version for Azure OpenAI
openai.api_key = AZURE_OPENAI_API_KEY
openai.api_base = AZURE_OPENAI_ENDPOINT  # Set to your Azure OpenAI endpoint


# Check if all environment variables are loaded correctly
if not all([BLOB_CONNECTION_STRING, AZURE_OPENAI_API_KEY, AZURE_SEARCH_ENDPOINT, AZURE_SEARCH_API_KEY, AZURE_SEARCH_INDEX_NAME, SYSTEM_MESSAGE, AZURE_OPENAI_ENDPOINT]):
    raise ValueError("One or more environment variables are missing.")

# Initialize BlobServiceClient globally
blob_service_client = BlobServiceClient.from_connection_string(BLOB_CONNECTION_STRING)

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS) #APIs or services where you want to allow anyone to send requests without needing authentication

# IndexDocuments function: Handles the indexing of documents
@app.route(route="IndexDocuments", methods=[ "GET","POST"]) #decarator(requst)
def IndexDocuments(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Processing IndexDocuments request.')


    if req.method == "GET":
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": "Failed to read file from blob storage"}),
            mimetype="application/json",
            status_code=400  # Bad Request
        )
    
    try:
        # Step 1: Parse incoming JSON body
        req_body = req.get_json()
        logging.info(f"Received request body: {json.dumps(req_body)}") # incoing json data

        doc_link = req_body.get("doc_link")  # extract the docuent link 
        #check for the presence 
        if not doc_link:
            logging.error("doc_link is required but not provided.")
            return func.HttpResponse(
                json.dumps({"status": "FAILED", "error": "doc_link is required"}),#Python object to JSON string.
                mimetype="application/json", # type of data being sent 
                status_code=400
            )

        logging.info(f"Document link: {doc_link}")

        # Step 2: Validate doc_link (check if it's a valid URL and ends with .pdf or .docx)
        if not validate_blob_link(doc_link):
            logging.error("Invalid blob storage link or unsupported file format.")
            return func.HttpResponse(
                json.dumps({"status": "FAILED", "error": "Invalid blob storage link or unsupported file format"}),
                mimetype="application/json",
                status_code=400 # Bad Request
            )

        # Step 3: Read the content from the blob storage
        file_content = read_blob_content(doc_link)
        if not file_content:
            logging.error("Failed to read file content from blob storage.")
            return func.HttpResponse(
                json.dumps({"status": "FAILED", "error": "Failed to read file from blob storage"}),
                mimetype="application/json",
                status_code=400
            )

        # Step 4: Split the content into chunks
        document_chunks = split_into_chunks(file_content)
        logging.info(f"Document split into {len(document_chunks)} chunks.")

        # Step 5: Create embeddings for each chunk using OpenAI embeddings
        embeddings = create_embeddings(document_chunks)
        logging.info(f"Generated {len(embeddings)} embeddings.")

        # Step 6: Store embeddings in Azure Cognitive Search
        store_embeddings_in_search_index(embeddings, document_chunks)
        logging.info(f"Stored {len(document_chunks)} document chunks in Azure Cognitive Search.")

        return func.HttpResponse(
            json.dumps({"status": "COMPLETED", "error": None}),
            mimetype="application/json",
            status_code=200
        )

    except Exception as e:
        logging.error(f"Error in IndexDocuments function: {str(e)}")
        return func.HttpResponse(
            json.dumps({"status": "FAILED", "error": f"Error: {str(e)}"}),
            mimetype="application/json", # output formate
            status_code=500
        )


# Function to validate the document URL
def validate_blob_link(doc_link):
    """Validate the blob storage link and file format."""
    blob_pattern = r"^https?://.*\.pdf$|^https?://.*\.docx$"
    return bool(re.match(blob_pattern, doc_link))


# Function to read content from blob storage
def read_blob_content(doc_link):
    """Read content from blob storage."""
    try:
        # Extract container name and blob name from the URL
        match = re.match(r"https://([a-z0-9]+)\.blob\.core\.windows\.net/([a-zA-Z0-9-]+)/(.+)", doc_link)
        if not match:
            logging.error(f"Invalid blob URL format: {doc_link}")
            return None

        container_name = match.group(2) # Container name
        blob_name = match.group(3) # Blob name

        logging.info(f"Container name: {container_name}, Blob name: {blob_name}")

        # Access the blob client
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)

        # Download the blob content
        blob_data = blob_client.download_blob() #cloud
        file_content = blob_data.readall() #reads the file to get its content

        # Log the size of the file downloaded
        logging.info(f"Downloaded {len(file_content)} bytes from blob.")

        # Process the file based on its extension (PDF/DOCX)
        if doc_link.endswith('.pdf'):
            return extract_pdf_text(file_content)
        elif doc_link.endswith('.docx'):
            return extract_docx_text(file_content)
        else:
            logging.error(f"Unsupported file type. Only PDF and DOCX are supported: {doc_link}")
            return None

    except Exception as e:
        logging.error(f"Error reading blob content: {str(e)}")
        return None

# Function to extract text from PDF
def extract_pdf_text(file_content):
    """Extract text from a PDF file."""
    try:
        # Convert bytes to a file-like object
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting PDF text: {str(e)}")
        return None


# Function to extract text from DOCX
def extract_docx_text(file_content):
    """Extract text from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {str(e)}")
        return None


# Function to split content into chunks with improved formatting and structure
def split_into_chunks(content):
    """Split document content into chunks, preserving paragraphs and sentence boundaries."""
    try:
        # First, split the content into paragraphs (separate by double newlines)
        paragraphs = content.split('\n\n')
        
        chunks = []
        for para in paragraphs:
            # Split each paragraph into sentences (this will preserve structure)
            sentences = sent_tokenize(para)
            
            # Add the sentences to chunks, ensuring we're splitting reasonably
            current_chunk = ""
            for sentence in sentences:
                if len(current_chunk) + len(sentence) > 1000:  # max chunk length for indexing (example: 1000 chars)
                    chunks.append(current_chunk)
                    current_chunk = sentence  # Start a new chunk with the current sentence
                else:
                    current_chunk += " " + sentence
            
            # Add any remaining chunk
            if current_chunk:
                chunks.append(current_chunk)
        
        logging.info(f"Document split into {len(chunks)} chunks.")
        return chunks
    except Exception as e:
        logging.error(f"Error in split_into_chunks: {str(e)}")
        return []


def create_embeddings(chunks):
    """Create embeddings for each chunk using Azure OpenAI embeddings."""
    embeddings = []
    AZURE_OPENAI_API_KEY = "d57bdf8c8fd54d98aea4f20e0a183479"
    AZURE_OPENAI_API_BASE = "https://openai-azuretesting.openai.azure.com/"
    AZURE_OPENAI_API_VERSION = "2024-05-01-preview"
    embedding_model = AzureOpenAIEmbeddings(
        model="text-embedding-3-small",  # Use a model suitable for embeddings
        azure_endpoint=AZURE_OPENAI_API_BASE,
        api_key=AZURE_OPENAI_API_KEY,
        openai_api_version=AZURE_OPENAI_API_VERSION
    )

    for chunk in chunks:
        try:
            if chunk.strip() == "":  # Skip empty chunks
                continue

            # Generate embedding for the chunk
            embedding = embedding_model.embed_query(chunk.strip())
            
            if embedding:
                embeddings.append(embedding)
            else:
                logging.error(f"No embedding returned for chunk: {chunk}")
                continue

        except Exception as e:
            logging.error(f"Error generating embedding for chunk: {str(e)}")
            continue  # Skip the chunk on error and proceed with the next one

    if len(embeddings) == 0:
        logging.error("No valid embeddings were generated.")
        return []  # Return an empty list instead of None to avoid len() errors

    return embeddings


# Define embedding function converts text into embeddings semantic search, content recommendation
def embedding_function(text):
    """Generate embeddings for the given text using OpenAI embeddings."""
    AZURE_OPENAI_API_KEY = "d57bdf8c8fd54d98aea4f20e0a183479"
    AZURE_OPENAI_API_BASE = "https://openai-azuretesting.openai.azure.com/"
    AZURE_OPENAI_API_VERSION = "2024-05-01-preview"
    embedding_model = AzureOpenAIEmbeddings(
        model="text-embedding-3-small",  # Use a model suitable for embeddings
        azure_endpoint=AZURE_OPENAI_API_BASE,
        api_key=AZURE_OPENAI_API_KEY,
        openai_api_version=AZURE_OPENAI_API_VERSION
    )
    return embedding_model.embed_query(text)

# Function to store the chunks with embeddings in Azure Cognitive Search
def store_embeddings_in_search_index(embeddings, document_chunks):
    """Store the document chunks and embeddings in Azure Cognitive Search."""
    
    # Initialize the AzureSearch client with the required parameters
    search_client = AzureSearch(
        azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
        azure_search_key=AZURE_SEARCH_API_KEY,
        embedding_function=embedding_function,  # Pass the embedding function
        index_name=AZURE_SEARCH_INDEX_NAME
    )
    
    # Index the document chunks with their embeddings as separate documents
    batch_documents = []
    # easily access both the index and value
    for i, chunk in enumerate(document_chunks):
        embedding = embeddings[i]
        
        # Ensure the embedding is a list/array of floats (dimensions: 1536)
        if not isinstance(embedding, list):
            embedding = [embedding]  # Wrap the embedding in a list if it's a single value
        
        if len(embedding) != 1536:  # Check if the embedding has the expected size
            logging.error(f"Embedding size mismatch for document {i}. Expected 1536, got {len(embedding)}")
            continue  # Skip this document if the embedding is invalid
        
        # Create a document dictionary with the required fields
        document = CustomDocument(
                doc_id=str(i),  # Assign a unique ID for each chunk
                content=chunk,
                embedding=embedding,  # Ensure embedding is properly associated
                metadata={}  # Include empty metadata if not provided
            )
        
        logging.info(f"Document {i}: {document}")  # Log the document before indexing
        batch_documents.append(document)  # Add document to batch for indexing
    
    # Add documents in bulk to the Azure Cognitive Search index
    if batch_documents:
        try:
            search_client.add_documents(batch_documents)  # This will index the documents
            logging.info(f"Indexed {len(batch_documents)} chunks into Azure Search.")
        except Exception as e:
            logging.error(f"Error while indexing documents: {str(e)}")
    else:
        logging.warning("No valid documents to index.")

  # search_client.add_documents(document_chunks)

# QueryKnowledgeBase function: Handles querying the knowledge base
@app.route(route="QueryKnowledgeBase", methods=["GET"])
def QueryKnowledgeBase(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Processing QueryKnowledgeBase request.')

    try:
        query = req.params.get('query')
        index_name = req.params.get('index_name')

        if not query or not index_name:
            return func.HttpResponse(
                json.dumps({"response": None, "error": "Both 'query' and 'index_name' are required"}),
                mimetype="application/json",
                status_code=400
            )

        # Generate embedding for the query using Azure OpenAI embeddings
        AZURE_OPENAI_API_KEY = "d57bdf8c8fd54d98aea4f20e0a183479"
        AZURE_OPENAI_API_BASE = "https://openai-azuretesting.openai.azure.com/"
        AZURE_OPENAI_API_VERSION = "2024-05-01-preview"
        
        embedding_model = AzureOpenAIEmbeddings(
            model="text-embedding-3-small",  # Use a model suitable for embeddings
            azure_endpoint=AZURE_OPENAI_API_BASE,
            api_key=AZURE_OPENAI_API_KEY,
            openai_api_version=AZURE_OPENAI_API_VERSION
        )

        # Search for the query embedding in the Azure Cognitive Search index
        search_client = AzureSearch(
            azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
            azure_search_key=AZURE_SEARCH_API_KEY,
            index_name=index_name,
            embedding_function=embedding_model
        )

        # Perform vector search and get the top results with "similarity" search type
        results = search_client.search(query, search_type="similarity") #for matching 

        # Use page_content or the appropriate attribute for the content in the result
        response_content = [result.page_content for result in results]

        if response_content:
            # Construct the system message from response content
            system_message = SYSTEM_MESSAGE + "\n" + "\n".join(response_content)

            # Query the GPT model with the constructed prompt using the new API interface
            openai.api_key = AZURE_OPENAI_API_KEY
            openai.api_base = AZURE_OPENAI_API_BASE  # Set your Azure endpoint here

            # Properly format the messages as an array of dictionaries
            messages = [
                {"role": "system", "content": system_message},
                {"role": "user", "content": query}
            ]

            # Use openai.ChatCompletion.create for the new API
            gpt_response = embedding_client.chat.completions.create(
                model="gpt-35-turbo-16k",
                messages=messages,  # Pass messages as an array
                temperature=0.25
            )

            response = gpt_response.choices[0].message.content
        else:
            response = None

        return func.HttpResponse(
            json.dumps({"response": response, "error": None}),
            mimetype="application/json",
            status_code=200
        )

    except Exception as e:
        logging.error(f"Error processing query: {str(e)}")
        return func.HttpResponse(
            json.dumps({"response": None, "error": f"Error: {str(e)}"}),
            mimetype="application/json",
            status_code=500
        )


embedding_client = AzureOpenAI(
    azure_endpoint="https://openai-azuretesting.openai.azure.com/",
    api_key="d57bdf8c8fd54d98aea4f20e0a183479",
    api_version="2024-02-15-preview"
)

class CustomDocument:
    def __init__(self, doc_id, content, embedding=None, metadata=None):
        self.id = doc_id
        self.page_content = content  # Use 'page_content' instead of 'content' to match expected format
        self.embedding = embedding
        self.metadata = metadata or {}  # Default to an empty dictionary if metadata is not provided
 
    def to_dict(self):# method converts the object into a dictionary format
        return {
            "id": self.id,
            "page_content": self.page_content,  # Ensure the correct field is used for the content
            "embedding": self.embedding,
            "metadata": self.metadata  # Include metadata field in the dictionary
        }
    